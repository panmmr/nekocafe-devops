#!/usr/bin/env python3
"""
自动填充实验三 docx 模板
将 docs/*.md 中的内容填入任务书提供的 docx 模板
"""
import os
import sys

TEMPLATE_DIR = r"C:\Users\33492\Desktop\软件工程实验任务书_2026春\实验三_DevOps流水线与容器化部署\产出模板"
OUTPUT_DIR = r"C:\Users\33492\Desktop\计算机23-2_230502209_潘美儒_实验三_v1.0\filled_templates"

STUDENT_INFO = {
    "班级": "计算机23-2",
    "学号": "230502209",
    "姓名": "潘美儒",
    "班级 / 学号 / 姓名": "计算机23-2 / 230502209 / 潘美儒",
    "___________ / ___________ / ___________": "计算机23-2 / 230502209 / 潘美儒",
    "提交日期": "2026-06-15",
    "20___ - ___ - ___": "2026-06-15",
}


def replace_in_paragraphs(doc, replacements: dict):
    """Replace text in all paragraphs of the document"""
    from docx.oxml.ns import qn
    for para in doc.paragraphs:
        for run in para.runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for old, new in replacements.items():
                            if old in run.text:
                                run.text = run.text.replace(old, new)


def fill_d3_1():
    """Fill D3-1 DevOps设计方案模板"""
    from docx import Document
    template = os.path.join(TEMPLATE_DIR, "D3-1_DevOps设计方案_模板.docx")
    if not os.path.exists(template):
        print(f"  SKIP: Template not found at {template}")
        return

    doc = Document(template)
    replacements = {
        **STUDENT_INFO,
        "（请在此处填写……）": "（详见 docs/D3-1_DevOps设计方案.md 完整内容）",
    }
    replace_in_paragraphs(doc, replacements)

    out = os.path.join(OUTPUT_DIR, "D3-1_DevOps设计方案_潘美儒.docx")
    doc.save(out)
    print(f"  ✅ D3-1 → {out}")


def fill_d3_4():
    """Fill D3-4 CI/CD配置与运行截图模板"""
    from docx import Document
    template = os.path.join(TEMPLATE_DIR, "D3-4_CICD配置与运行截图_模板.docx")
    if not os.path.exists(template):
        print(f"  SKIP: Template not found at {template}")
        return

    doc = Document(template)
    replacements = {
        **STUDENT_INFO,
        "（请在此处填写本节正文……）": "（详见 GitHub Actions 运行截图和 docs/D3-4_CICD配置与运行截图.md）",
        "（粘贴 .github/workflows/ci.yml 完整内容）": "（见项目根目录 .github/workflows/ci.yml）",
        "（GitHub Environments + required reviewers……）": "dev: 无需审批 / staging: 1 reviewer / prod: 2 reviewers",
    }
    replace_in_paragraphs(doc, replacements)

    out = os.path.join(OUTPUT_DIR, "D3-4_CICD配置与运行截图_潘美儒.docx")
    doc.save(out)
    print(f"  ✅ D3-4 → {out}")


def fill_d3_6():
    """Fill D3-6 可观测性配置与Dashboard截图模板"""
    from docx import Document
    template = os.path.join(TEMPLATE_DIR, "D3-6_可观测性配置与Dashboard截图_模板.docx")
    if not os.path.exists(template):
        print(f"  SKIP: Template not found at {template}")
        return

    doc = Document(template)
    replacements = {
        **STUDENT_INFO,
        "（请在此处填写本节正文……）": "（详见 docs/D3-6_可观测性配置与Dashboard截图.md 和 Grafana 截图）",
        "（语言 SDK + 自动注入还是手动埋点）": "Python OpenTelemetry SDK 1.28.2，FastAPIInstrumentor 自动注入 + trace_id 中间件手动注入",
        "（HEAD / TAIL / PROBABILITY 与采样率）": "开发环境 ALWAYS_ON (100%)，生产建议 TAIL sampling",
        "（time / level / traceId / spanId / service……）": "timestamp / level / service / message / module / traceId",
        "（手机号 / 身份证 / 邮箱 的处理）": "手机号脱敏为 138****8000，身份证不记录，邮箱 p***@neko.cafe",
        "（Rate / Errors / Duration）": "rate(http_requests_total[1m]) / rate(http_requests_total{status~5..}[5m]) / histogram_quantile(0.99, ...)",
        "（Utilization / Saturation / Errors）": "container_cpu_usage_seconds_total / container_memory_working_set_bytes",
        "（按\"业务 / 应用 / 中间件 / 基础设施\"四层展示）": "业务(QPS/预订成功率) → 应用(P99延迟/错误率) → 中间件(DB/Redis连接数) → 基础设施(CPU/内存)",
        "（粘贴 PrometheusRule YAML……）": "（见 k8s/monitoring/prometheus-rules/alerts.yaml，共 5 条规则）",
        "（钉钉 / 飞书 / 短信 / 电话升级策略）": "staging: 钉钉机器人 / prod: 钉钉 + 短信 + 电话升级",
        "（描述演练范围、注入工具、观测结果、改进项）": "N/A — Chaos Mesh 为可选加分项，本次实验暂未实施",
    }
    replace_in_paragraphs(doc, replacements)

    out = os.path.join(OUTPUT_DIR, "D3-6_可观测性配置与Dashboard截图_潘美儒.docx")
    doc.save(out)
    print(f"  ✅ D3-6 → {out}")


def fill_d3_8():
    """Fill D3-8 演示视频脚本模板"""
    from docx import Document
    template = os.path.join(TEMPLATE_DIR, "D3-8_演示视频脚本_模板.docx")
    if not os.path.exists(template):
        print(f"  SKIP: Template not found at {template}")
        return

    doc = Document(template)
    replacements = {
        **STUDENT_INFO,
        "（请在此处填写……）": "（详见 docs/D3-8_演示视频脚本.md — 完整 4 分 30 秒脚本，含时间段、画面内容、旁白）",
    }
    replace_in_paragraphs(doc, replacements)

    out = os.path.join(OUTPUT_DIR, "D3-8_演示视频脚本_潘美儒.docx")
    doc.save(out)
    print(f"  ✅ D3-8 → {out}")


def fill_d3_7():
    """Convert D3-7 markdown into a simple xlsx using openpyxl"""
    try:
        from openpyxl import Workbook
    except ImportError:
        print("  SKIP D3-7: openpyxl not installed. Run: pip install openpyxl")
        return

    wb = Workbook()
    ws = wb.active
    ws.title = "DORA指标报告"

    rows = [
        ["DORA 指标报告 — NekoCafé"],
        [""],
        ["指标", "数据源", "采集方法", "目标", "当前值"],
        ["部署频率 (DF)", "GitHub Actions 运行记录", "gh run list --workflow=cd.yml", "≥ 1次/天", "[运行后填写]"],
        ["变更前置时间 (LTTC)", "Git + GitHub Actions", "commit_timestamp → deploy_timestamp", "< 1小时", "[运行后填写]"],
        ["变更失败率 (CFR)", "GitHub Actions + Rollback", "回滚次数 / 总部署次数", "< 5%", "[运行后填写]"],
        ["恢复时间 (MTTR)", "Rollback 脚本执行日志", "rollback_start → rollback_end", "< 10分钟", "[运行后填写]"],
        [""],
        ["班级", "计算机23-2"],
        ["学号", "230502209"],
        ["姓名", "潘美儒"],
        ["日期", "2026-06-15"],
    ]
    for row in rows:
        ws.append(row)

    out = os.path.join(OUTPUT_DIR, "D3-7_DORA指标报告_潘美儒.xlsx")
    wb.save(out)
    print(f"  ✅ D3-7 → {out}")


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    print("=== 自动填充实验三 docx 模板 ===\n")
    print(f"模板目录：{TEMPLATE_DIR}")
    print(f"输出目录：{OUTPUT_DIR}")
    print()

    if not os.path.exists(TEMPLATE_DIR):
        print("ERROR: 模板目录不存在！请确认路径：")
        print(f"  {TEMPLATE_DIR}")
        sys.exit(1)

    fill_d3_1()
    fill_d3_4()
    fill_d3_6()
    fill_d3_7()
    fill_d3_8()

    print(f"\n完成！已填充的模板保存在：{OUTPUT_DIR}/")
    print()
    print("注意事项：")
    print("  - D3-2（源代码仓库）：本身就是代码仓库，无需单独文档")
    print("  - D3-3（Dockerfile与镜像扫描）：请手动粘贴 Trivy 扫描结果和镜像大小")
    print("  - D3-5（K8s部署清单）：请手动粘贴 kubectl 输出截图")
    print("  - D3-9（答辩PPT）：需手动在 PowerPoint 中制作（见 D3-9_答辩PPT大纲.md）")


if __name__ == "__main__":
    main()
